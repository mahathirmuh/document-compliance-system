"""Generated, non-confidential fixtures for the pure Phase 6 extractors."""

from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import pymupdf
import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.worksheet.table import Table

from app.schemas.extraction import (
    ExtractedBlockType,
    ExtractionResultStatus,
)
from app.services.extraction.base_extractor import (
    ExtractionCancelledError,
    ExtractionError,
    ExtractionResourceLimitError,
    UnsupportedExtractionFormatError,
)
from app.services.extraction.docx.docx_extractor import DOCXExtractor
from app.services.extraction.extractor_factory import ExtractorFactory
from app.services.extraction.pdf.pdf_extractor import PDFExtractor
from app.services.extraction.text_normalizer import (
    calculate_content_hash,
    count_characters,
    count_words,
    normalize_text,
)
from app.services.extraction.xlsx.xlsx_extractor import XLSXExtractor


def _write_pdf(
    path: Path,
    pages: list[str | None],
    *,
    password: str | None = None,
) -> None:
    document = pymupdf.open()
    for text in pages:
        page = document.new_page()
        if text is None:
            pixmap = pymupdf.Pixmap(
                pymupdf.csRGB,
                pymupdf.IRect(0, 0, 100, 100),
                False,
            )
            pixmap.clear_with(255)
            page.insert_image(page.rect, stream=pixmap.tobytes("png"))
        elif text:
            page.insert_text((72, 72), text)
    save_options = {}
    if password is not None:
        save_options = {
            "encryption": pymupdf.PDF_ENCRYPT_AES_256,
            "owner_pw": "generated-owner-password",
            "user_pw": password,
        }
    document.save(path, **save_options)
    document.close()


def _add_external_hyperlink(paragraph, text: str) -> None:
    relationship_id = paragraph.part.relate_to(
        "https://example.invalid/not-fetched",
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _inject_unsafe_doctype(path: Path, member_name: str) -> None:
    replacement_path = path.with_suffix(f"{path.suffix}.replacement")
    with ZipFile(path) as source:
        members = {
            item.filename: source.read(item.filename)
            for item in source.infolist()
        }
    xml = members[member_name]
    declaration = (
        b'<!DOCTYPE root [<!ENTITY generated SYSTEM "file:///never-read">]>'
    )
    insertion_point = xml.find(b"?>")
    if insertion_point >= 0:
        insertion_point += 2
        xml = xml[:insertion_point] + declaration + xml[insertion_point:]
    else:
        xml = declaration + xml
    members[member_name] = xml
    with ZipFile(replacement_path, "w", ZIP_DEFLATED) as destination:
        for name, payload in members.items():
            destination.writestr(name, payload)
    replacement_path.replace(path)


@pytest.mark.parametrize(
    ("extension", "expected_type"),
    [
        ("pdf", PDFExtractor),
        (".DOCX", DOCXExtractor),
        (" xlsx ", XLSXExtractor),
    ],
)
def test_extractor_factory_selects_only_supported_formats(
    extension: str,
    expected_type: type,
) -> None:
    assert isinstance(
        ExtractorFactory.get_extractor(extension),
        expected_type,
    )


def test_extractor_factory_rejects_unsupported_format() -> None:
    with pytest.raises(UnsupportedExtractionFormatError) as error:
        ExtractorFactory.get_extractor("pptx")
    assert error.value.code == "UNSUPPORTED_EXTRACTION_FORMAT"


@pytest.mark.asyncio
async def test_pdf_extracts_pages_text_blocks_and_bounding_boxes(
    tmp_path: Path,
) -> None:
    source = tmp_path / "multi-page.pdf"
    _write_pdf(
        source,
        [
            "Document Control Procedure with selectable text",
            "Second page with selectable text",
        ],
    )

    result = await PDFExtractor().extract(source, {})

    assert result.status is ExtractionResultStatus.COMPLETED
    assert result.metadata["totalPages"] == 2
    assert len(result.containers) == 2
    first_block = result.containers[0].blocks[0]
    assert first_block.block_type is ExtractedBlockType.TEXT
    assert first_block.source_reference == "PDF:page=1:block=1"
    assert len(first_block.location["bbox"]) == 4
    assert first_block.location["width"] > 0
    assert result.requires_ocr is False


@pytest.mark.asyncio
async def test_pdf_detects_scan_and_partial_scan(tmp_path: Path) -> None:
    scanned_source = tmp_path / "scanned.pdf"
    partial_source = tmp_path / "partial-scan.pdf"
    _write_pdf(scanned_source, [None, None])
    _write_pdf(
        partial_source,
        ["This page has enough selectable text for extraction.", None],
    )

    scanned = await PDFExtractor().extract(scanned_source, {})
    partial = await PDFExtractor().extract(partial_source, {})

    assert scanned.status is ExtractionResultStatus.OCR_REQUIRED
    assert scanned.requires_ocr is True
    assert scanned.has_selectable_text is False
    assert partial.status is ExtractionResultStatus.PARTIALLY_COMPLETED
    assert partial.requires_ocr is True
    assert partial.has_selectable_text is True


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("fixture_name", "expected_code"),
    [
        ("blank", "PDF_EMPTY"),
        ("corrupt", "PDF_CORRUPT"),
        ("protected", "PDF_PASSWORD_REQUIRED"),
    ],
)
async def test_pdf_rejects_empty_corrupt_and_password_protected_files(
    tmp_path: Path,
    fixture_name: str,
    expected_code: str,
) -> None:
    source = tmp_path / f"{fixture_name}.pdf"
    if fixture_name == "blank":
        _write_pdf(source, [""])
    elif fixture_name == "protected":
        _write_pdf(source, ["Protected text"], password="secret")
    else:
        source.write_bytes(b"%PDF-corrupt-generated-fixture")

    with pytest.raises(ExtractionError) as error:
        await PDFExtractor().extract(source, {})
    assert error.value.code == expected_code


@pytest.mark.asyncio
async def test_pdf_enforces_page_limit_and_cancellation(tmp_path: Path) -> None:
    source = tmp_path / "bounded.pdf"
    _write_pdf(source, ["Page one content", "Page two content"])

    with pytest.raises(ExtractionResourceLimitError):
        await PDFExtractor().extract(source, {"pdf_max_pages": 1})

    with pytest.raises(ExtractionCancelledError):
        await PDFExtractor().extract(
            source,
            {"cancellation_checker": lambda: True},
        )


@pytest.mark.asyncio
async def test_docx_preserves_body_order_and_extracts_structures(
    tmp_path: Path,
) -> None:
    source = tmp_path / "ordered.docx"
    document = Document()
    document.add_heading("Document Heading", level=2)
    document.add_paragraph("Paragraph before table")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).merge(table.cell(1, 1)).text = "Merged value"
    document.add_paragraph("Paragraph after table")
    document.sections[0].header.paragraphs[0].text = "Controlled header"
    document.sections[0].footer.paragraphs[0].text = "Controlled footer"
    document.save(source)

    result = await DOCXExtractor().extract(source, {})

    body = result.containers[0]
    block_types = [block.block_type for block in body.blocks]
    assert block_types[:3] == [
        ExtractedBlockType.HEADING,
        ExtractedBlockType.PARAGRAPH,
        ExtractedBlockType.TABLE,
    ]
    table_index = block_types.index(ExtractedBlockType.TABLE)
    trailing_paragraph_index = max(
        index
        for index, block_type in enumerate(block_types)
        if block_type is ExtractedBlockType.PARAGRAPH
    )
    assert table_index < trailing_paragraph_index
    assert body.blocks[0].heading_level == 2
    assert body.tables[0].cells[-1].column_span == 2
    assert any(
        container.container_type.value == "DOCX_HEADER"
        and "Controlled header" in container.raw_text
        for container in result.containers
    )

    with pytest.raises(ExtractionResourceLimitError):
        await DOCXExtractor().extract(
            source,
            {"docx_max_table_cells": 2},
        )
    assert any(
        container.container_type.value == "DOCX_FOOTER"
        and "Controlled footer" in container.raw_text
        for container in result.containers
    )


@pytest.mark.asyncio
async def test_docx_reads_hyperlink_text_without_following_relationship(
    tmp_path: Path,
) -> None:
    source = tmp_path / "external-link.docx"
    document = Document()
    paragraph = document.add_paragraph("Before ")
    _add_external_hyperlink(paragraph, "linked text")
    document.save(source)

    inspected = await DOCXExtractor().inspect(source)
    result = await DOCXExtractor().extract(source, {})

    assert inspected["externalRelationshipCount"] == 1
    assert result.metadata["externalRelationshipsFollowed"] is False
    assert "linked text" in result.containers[0].raw_text


@pytest.mark.asyncio
async def test_docx_reports_unsupported_content_and_limits(
    tmp_path: Path,
) -> None:
    source = tmp_path / "unsupported.docx"
    document = Document()
    paragraph = document.add_paragraph("Visible text")
    paragraph._p.append(OxmlElement("w:drawing"))
    document.add_paragraph("Second paragraph")
    document.save(source)

    result = await DOCXExtractor().extract(source, {})
    assert result.status is ExtractionResultStatus.PARTIALLY_COMPLETED
    assert any("Drawing" in warning for warning in result.warnings)

    with pytest.raises(ExtractionResourceLimitError):
        await DOCXExtractor().extract(
            source,
            {"docx_max_paragraphs": 1},
        )


@pytest.mark.asyncio
@pytest.mark.parametrize("fixture_name", ["empty", "corrupt"])
async def test_docx_handles_empty_and_corrupt_files(
    tmp_path: Path,
    fixture_name: str,
) -> None:
    source = tmp_path / f"{fixture_name}.docx"
    if fixture_name == "empty":
        Document().save(source)
        expected_code = "DOCX_EMPTY"
    else:
        source.write_bytes(b"generated-corrupt-docx")
        expected_code = "DOCX_CORRUPT"

    with pytest.raises(ExtractionError) as error:
        await DOCXExtractor().extract(source, {})
    assert error.value.code == expected_code


@pytest.mark.asyncio
async def test_docx_rejects_dtd_and_entity_declarations(tmp_path: Path) -> None:
    source = tmp_path / "unsafe-xml.docx"
    document = Document()
    document.add_paragraph("This content must never trigger entity resolution.")
    document.save(source)
    _inject_unsafe_doctype(source, "word/document.xml")

    with pytest.raises(ExtractionError) as error:
        await DOCXExtractor().extract(source, {})
    assert error.value.code == "DOCX_CORRUPT"


@pytest.mark.asyncio
async def test_xlsx_extracts_sheets_cells_formulas_merges_and_visibility(
    tmp_path: Path,
) -> None:
    source = tmp_path / "workbook.xlsx"
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Register"
    worksheet["A1"] = "Document Code"
    worksheet["B1"] = "Code Length"
    worksheet["A2"] = "MTI-HRM-POL-001"
    worksheet["B2"] = "=LEN(A2)"
    worksheet["A2"].hyperlink = "https://example.invalid/not-fetched"
    worksheet.merge_cells("A4:C4")
    worksheet["A4"] = "Merged heading"
    worksheet.merge_cells("A6:B6")
    worksheet.freeze_panes = "A2"
    worksheet.row_dimensions[2].hidden = True
    worksheet.column_dimensions["B"].hidden = True
    worksheet.protection.sheet = True
    worksheet.add_table(Table(displayName="RegisterTable", ref="A1:B2"))
    hidden = workbook.create_sheet("Hidden")
    hidden.sheet_state = "hidden"
    hidden["A1"] = "Retained hidden content"
    workbook.save(source)

    result = await XLSXExtractor().extract(source, {})

    assert result.status is ExtractionResultStatus.COMPLETED
    assert len(result.containers) == 2
    blocks = result.containers[0].blocks
    formula = next(
        block
        for block in blocks
        if block.block_type is ExtractedBlockType.FORMULA
    )
    merged = next(
        block
        for block in blocks
        if block.block_type is ExtractedBlockType.MERGED_CELL
    )
    assert formula.text == "=LEN(A2)"
    assert formula.metadata["formulaExecuted"] is False
    assert formula.metadata["cachedValue"] is None
    assert formula.metadata["columnHidden"] is True
    linked_cell = next(block for block in blocks if block.location["row"] == 2)
    assert linked_cell.metadata["rowHidden"] is True
    assert linked_cell.metadata["hyperlink"].startswith("https://example.invalid")
    assert linked_cell.metadata["hyperlinkFollowed"] is False
    assert merged.metadata["range"] == "A4:C4"
    assert merged.metadata["columnSpan"] == 3
    assert any(
        block.metadata.get("range") == "A6:B6" and block.text == ""
        for block in blocks
    )
    assert result.containers[0].metadata["freezePane"] == "A2"
    assert result.containers[0].metadata["tableNames"] == ["RegisterTable"]
    assert result.containers[0].metadata["sheetProtected"] is True
    assert result.containers[1].metadata["sheetState"] == "hidden"


@pytest.mark.asyncio
async def test_xlsx_enforces_sheet_cell_formula_and_row_limits(
    tmp_path: Path,
) -> None:
    source = tmp_path / "limited.xlsx"
    workbook = Workbook()
    worksheet = workbook.active
    worksheet["A1"] = "one"
    worksheet["A2"] = "two"
    worksheet["A3"] = "=1+1"
    workbook.create_sheet("Second")["A1"] = "second"
    workbook.save(source)

    cases = [
        ({"xlsx_max_worksheets": 1}, "XLSX_TOO_MANY_SHEETS"),
        ({"xlsx_max_cells_per_workbook": 1}, "XLSX_TOO_MANY_CELLS"),
        ({"xlsx_max_formulas": 1, "xlsx_max_cells_per_workbook": 10}, None),
        ({"xlsx_max_rows_per_sheet": 2}, "XLSX_WORKBOOK_TOO_LARGE"),
    ]
    for context, expected_code in cases:
        if expected_code is None:
            result = await XLSXExtractor().extract(source, context)
            assert result.metadata["totalFormulas"] == 1
            continue
        with pytest.raises(ExtractionResourceLimitError) as error:
            await XLSXExtractor().extract(source, context)
        assert error.value.code == expected_code


@pytest.mark.asyncio
@pytest.mark.parametrize("fixture_name", ["empty", "corrupt"])
async def test_xlsx_handles_empty_and_corrupt_workbooks(
    tmp_path: Path,
    fixture_name: str,
) -> None:
    source = tmp_path / f"{fixture_name}.xlsx"
    if fixture_name == "empty":
        Workbook().save(source)
        expected_code = "XLSX_EMPTY"
    else:
        source.write_bytes(b"generated-corrupt-xlsx")
        expected_code = "XLSX_CORRUPT"

    with pytest.raises(ExtractionError) as error:
        await XLSXExtractor().extract(source, {})
    assert error.value.code == expected_code


@pytest.mark.asyncio
async def test_xlsx_formula_limit_does_not_execute_formula(
    tmp_path: Path,
) -> None:
    source = tmp_path / "formulas.xlsx"
    workbook = Workbook()
    worksheet = workbook.active
    worksheet["A1"] = "=1+1"
    worksheet["A2"] = "=2+2"
    workbook.save(source)

    with pytest.raises(ExtractionResourceLimitError) as error:
        await XLSXExtractor().extract(source, {"xlsx_max_formulas": 1})
    assert error.value.code == "XLSX_WORKBOOK_TOO_LARGE"

    with pytest.raises(ExtractionResourceLimitError):
        await XLSXExtractor().extract(source, {"xlsx_max_formulas": 0})


@pytest.mark.asyncio
async def test_xlsx_rejects_dtd_and_entity_declarations(tmp_path: Path) -> None:
    source = tmp_path / "unsafe-xml.xlsx"
    workbook = Workbook()
    workbook.active["A1"] = "Safe source value"
    workbook.save(source)
    _inject_unsafe_doctype(source, "xl/worksheets/sheet1.xml")

    with pytest.raises(ExtractionError) as error:
        await XLSXExtractor().extract(source, {})
    assert error.value.code == "XLSX_CORRUPT"


@pytest.mark.asyncio
async def test_text_normalization_counts_and_hash_are_deterministic(
    tmp_path: Path,
) -> None:
    assert normalize_text("Cafe\u0301  \r\n中文。\x00\r\n") == "Café\n中文。"
    assert count_characters("A\x00B") == 2
    assert count_words("alpha  beta\n中文") == 3

    source = tmp_path / "hash.pdf"
    _write_pdf(source, ["hash fixture text"])
    first = await PDFExtractor().extract(source, {})
    second = await PDFExtractor().extract(source, {})
    assert calculate_content_hash(first.containers) == calculate_content_hash(
        second.containers
    )
    assert len(calculate_content_hash([])) == 64
