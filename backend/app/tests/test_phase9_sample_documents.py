"""Deterministic synthetic Phase 9 fixture regressions."""

from __future__ import annotations

from hashlib import sha256
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from openpyxl import load_workbook

from scripts.generate_phase9_sample_documents import (
    EXPECTED_RELATIVE_PATHS,
    FIXTURE_AUTHOR,
    FIXTURE_SUBJECT,
    generate,
)

MAX_FIXTURE_BYTES = 250_000


def _digest(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def _docx_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join((*paragraphs, *cells))


def test_phase9_fixture_generator_is_exact_bounded_and_deterministic(
    tmp_path: Path,
) -> None:
    first_root = tmp_path / "first"
    second_root = tmp_path / "second"
    first_outputs = generate(first_root)
    second_outputs = generate(second_root)

    assert tuple(path.relative_to(first_root) for path in first_outputs) == (
        EXPECTED_RELATIVE_PATHS
    )
    assert tuple(path.relative_to(second_root) for path in second_outputs) == (
        EXPECTED_RELATIVE_PATHS
    )
    assert [_digest(path) for path in first_outputs] == [
        _digest(path) for path in second_outputs
    ]

    for path in first_outputs:
        assert 1_000 < path.stat().st_size < MAX_FIXTURE_BYTES, path.name
        with ZipFile(path) as archive:
            assert archive.testzip() is None
            assert archive.namelist()


def test_phase9_docx_fixtures_are_readable_and_synthetic(tmp_path: Path) -> None:
    outputs = generate(tmp_path)

    for path in (item for item in outputs if item.suffix == ".docx"):
        document = Document(path)
        assert document.core_properties.author == FIXTURE_AUTHOR
        assert document.core_properties.subject == FIXTURE_SUBJECT
        assert document.tables
        text = _docx_text(path)
        assert "Bahasa Indonesia" in text
        assert "English" in text
        assert "\u4e2d\u6587" in text
        assert len(text) < 20_000


def test_phase9_xlsx_fixtures_are_readable_and_bounded(tmp_path: Path) -> None:
    outputs = generate(tmp_path)

    for path in (item for item in outputs if item.suffix == ".xlsx"):
        workbook = load_workbook(path, read_only=True, data_only=False)
        try:
            assert workbook.properties.creator == FIXTURE_AUTHOR
            assert workbook.properties.subject == FIXTURE_SUBJECT
            assert 1 <= len(workbook.sheetnames) <= 3
            for sheet in workbook.worksheets:
                assert 1 <= sheet.max_row <= 10
                assert 1 <= sheet.max_column <= 12
                assert any(
                    cell.value not in (None, "")
                    for row in sheet.iter_rows()
                    for cell in row
                )
        finally:
            workbook.close()


def test_phase9_fixtures_preserve_named_validation_evidence(
    tmp_path: Path,
) -> None:
    outputs = {
        path.relative_to(tmp_path).as_posix(): path for path in generate(tmp_path)
    }

    assert "24" in _docx_text(outputs["similarity/number-mismatch.docx"])
    assert "25" in _docx_text(outputs["similarity/number-mismatch.docx"])
    assert "15 Agustus 2026" in _docx_text(outputs["similarity/date-mismatch.docx"])
    assert "16 August 2026" in _docx_text(outputs["similarity/date-mismatch.docx"])
    assert "5 kg" in _docx_text(outputs["similarity/measurement-mismatch.docx"])
    assert "7 kg" in _docx_text(outputs["similarity/measurement-mismatch.docx"])
    assert "WI-TEST-009" in _docx_text(outputs["similarity/reference-mismatch.docx"])
    assert "WI-TEST-010" in _docx_text(outputs["similarity/reference-mismatch.docx"])
    assert "tidak boleh" in _docx_text(outputs["similarity/negation-mismatch.docx"])
    assert "may delete" in _docx_text(outputs["similarity/negation-mismatch.docx"])

    revision_one = _docx_text(outputs["revisions/revision-001.docx"])
    revision_two = _docx_text(outputs["revisions/revision-002.docx"])
    assert "REV-ITEM-002" in revision_one
    assert "REV-ITEM-002" not in revision_two
    assert "REV-ITEM-004" not in revision_one
    assert "REV-ITEM-004" in revision_two


def test_chinese_glossary_workbook_uses_canonical_import_headers(
    tmp_path: Path,
) -> None:
    workbook_path = next(
        path for path in generate(tmp_path) if path.name == "chinese-terms.xlsx"
    )
    workbook = load_workbook(workbook_path, read_only=True, data_only=False)
    try:
        assert workbook.sheetnames == ["Terms", "Translations", "Variants"]
        assert tuple(cell.value for cell in workbook["Terms"][1]) == (
            "profile_code",
            "term_code",
            "concept_name",
            "description",
            "term_type",
            "severity",
            "is_case_sensitive",
            "match_whole_word",
            "allow_inflection",
            "is_regex",
            "is_active",
            "notes",
        )
        assert tuple(cell.value for cell in workbook["Translations"][1]) == (
            "term_code",
            "language_code",
            "term_text",
            "is_preferred",
            "is_forbidden",
            "is_required",
            "priority",
            "usage_note",
            "example_text",
            "is_active",
        )
        assert tuple(cell.value for cell in workbook["Variants"][1]) == (
            "term_code",
            "language_code",
            "preferred_term_text",
            "variant_text",
            "variant_type",
            "is_allowed",
            "is_active",
        )
        translation_rows = list(
            workbook["Translations"].iter_rows(min_row=2, values_only=True)
        )
        assert {row[1] for row in translation_rows} == {"id", "en", "zh"}
        assert any(row[2] == "\u53d7\u63a7\u6587\u4ef6" for row in translation_rows)
    finally:
        workbook.close()
