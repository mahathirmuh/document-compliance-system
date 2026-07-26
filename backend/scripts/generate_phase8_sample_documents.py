"""Generate synthetic Phase 8 multilingual-compliance fixtures."""

from __future__ import annotations

import argparse
from pathlib import Path

import pymupdf
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font

ZH_PURPOSE = "\u76ee\u7684"
ZH_SCOPE = "\u8303\u56f4"
ZH_PROCEDURE = "\u7a0b\u5e8f"
ZH_PURPOSE_TEXT = (
    "\u672c\u7a0b\u5e8f\u7528\u4e8e\u89c4\u8303\u5408\u6210"
    "\u6587\u4ef6\u7684\u63a7\u5236\u548c\u5ba1\u67e5\u3002"
)
ZH_SCOPE_TEXT = (
    "\u672c\u8981\u6c42\u9002\u7528\u4e8e\u6240\u6709\u751f\u6210"
    "\u7684\u6d4b\u8bd5\u6587\u4ef6\u548c\u8bb0\u5f55\u3002"
)
ZH_PROCEDURE_TEXT = (
    "\u6587\u4ef6\u8d1f\u8d23\u4eba\u5fc5\u987b\u5ba1\u67e5"
    "\u5e76\u4fdd\u7559\u6bcf\u4e2a\u7248\u672c\u3002"
)


def _metadata(document: Document, title: str) -> None:
    document.core_properties.author = "Document Compliance System"
    document.core_properties.title = title
    document.core_properties.subject = (
        "Synthetic Phase 8 fixture without company content"
    )


def _add_trilingual_section(
    document: Document,
    *,
    heading: tuple[str, str, str],
    paragraphs: tuple[str, str, str],
    order: tuple[int, int, int] = (0, 1, 2),
) -> None:
    document.add_heading(
        " / ".join(heading),
        level=1,
    )
    for index in order:
        document.add_paragraph(paragraphs[index])


def _complete_docx(
    path: Path,
    *,
    include_chinese: bool = True,
    wrong_order: bool = False,
) -> None:
    document = Document()
    _metadata(document, path.stem)
    order = (1, 0, 2) if wrong_order else (0, 1, 2)
    sections = (
        (
            ("Tujuan", "Purpose", ZH_PURPOSE),
            (
                "Prosedur ini mengatur pengendalian dokumen sintetis.",
                "This procedure controls synthetic test documents.",
                ZH_PURPOSE_TEXT,
            ),
        ),
        (
            ("Ruang Lingkup", "Scope", ZH_SCOPE),
            (
                "Persyaratan ini berlaku untuk seluruh dokumen pengujian.",
                "These requirements apply to every generated test document.",
                ZH_SCOPE_TEXT,
            ),
        ),
        (
            ("Prosedur", "Procedure", ZH_PROCEDURE),
            (
                "Pemilik dokumen meninjau dan menyimpan setiap revisi.",
                "The document owner reviews and retains every revision.",
                ZH_PROCEDURE_TEXT,
            ),
        ),
    )
    for heading, paragraphs in sections:
        selected = paragraphs if include_chinese else (paragraphs[0], paragraphs[1], "")
        selected_order = tuple(
            index for index in order if include_chinese or index != 2
        )
        _add_trilingual_section(
            document,
            heading=heading,
            paragraphs=selected,
            order=selected_order,
        )
    document.save(path)


def _missing_purpose_docx(path: Path) -> None:
    document = Document()
    _metadata(document, path.stem)
    _add_trilingual_section(
        document,
        heading=("Ruang Lingkup", "Scope", ZH_SCOPE),
        paragraphs=(
            "Dokumen ini hanya berisi bagian ruang lingkup sintetis.",
            "This document only contains a synthetic scope section.",
            ZH_SCOPE_TEXT,
        ),
    )
    document.save(path)


def _incomplete_table_docx(path: Path) -> None:
    document = Document()
    _metadata(document, path.stem)
    document.add_heading("Prosedur / Procedure / " + ZH_PROCEDURE, level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ("Bahasa Indonesia", "English", "\u4e2d\u6587")
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        cell.text = value
    rows = (
        (
            "Periksa identitas dokumen.",
            "Check the document identity.",
            "\u68c0\u67e5\u6587\u4ef6\u6807\u8bc6\u3002",
        ),
        (
            "Simpan catatan persetujuan.",
            "Retain the approval record.",
            "",
        ),
    )
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values, strict=True):
            cell.text = value
    document.save(path)


def _workbook(path: Path, title: str) -> tuple[Workbook, object]:
    workbook = Workbook()
    workbook.properties.creator = "Document Compliance System"
    workbook.properties.title = title
    workbook.properties.subject = "Synthetic Phase 8 fixture without company content"
    sheet = workbook.active
    sheet.title = "Procedure"
    return workbook, sheet


def _columns_xlsx(path: Path, *, missing_cell: bool = False) -> None:
    workbook, sheet = _workbook(path, path.stem)
    sheet.append(("Bahasa Indonesia", "English", "\u4e2d\u6587"))
    sheet.append(
        (
            "Periksa kode dokumen sebelum digunakan.",
            "Check the document code before use.",
            (
                ""
                if missing_cell
                else "\u4f7f\u7528\u524d\u68c0\u67e5\u6587\u4ef6\u7f16\u7801\u3002"
            ),
        )
    )
    sheet.append(
        (
            "Simpan setiap catatan revisi.",
            "Retain every revision record.",
            "\u4fdd\u7559\u6bcf\u4efd\u4fee\u8ba2\u8bb0\u5f55\u3002",
        )
    )
    for cell in sheet[1]:
        cell.font = Font(bold=True)
    sheet.freeze_panes = "A2"
    workbook.save(path)


def _rows_xlsx(path: Path) -> None:
    workbook, sheet = _workbook(path, path.stem)
    sheet.append(("Language", "Text"))
    sheet.append(
        (
            "Bahasa Indonesia",
            "Pemilik dokumen melakukan pemeriksaan berkala.",
        )
    )
    sheet.append(
        (
            "English",
            "The document owner performs a periodic review.",
        )
    )
    sheet.append(
        (
            "\u4e2d\u6587",
            (
                "\u6587\u4ef6\u8d1f\u8d23\u4eba\u5b9a\u671f"
                "\u8fdb\u884c\u5ba1\u67e5\u3002"
            ),
        )
    )
    for cell in sheet[1]:
        cell.font = Font(bold=True)
    workbook.save(path)


def _pdf(
    path: Path,
    *,
    wrong_order: bool = False,
    low_confidence: bool = False,
) -> None:
    document = pymupdf.open()
    page = document.new_page(width=595, height=842)
    heading = "1. Tujuan / Purpose / " + ZH_PURPOSE
    page.insert_text((55, 70), heading, fontname="china-s", fontsize=17)
    lines = [
        (
            "Prosedur ini mengatur dokumen sintetis.",
            "helv",
        ),
        (
            "This procedure controls synthetic documents.",
            "helv",
        ),
        (
            ZH_PURPOSE_TEXT,
            "china-s",
        ),
    ]
    if wrong_order:
        lines = [lines[1], lines[0], lines[2]]
    y = 120
    for text, font_name in lines:
        page.insert_text(
            (55, y),
            text,
            fontname=font_name,
            fontsize=12,
        )
        y += 180 if low_confidence else 34
    page.insert_text(
        (55, min(y + 25, 760)),
        "2. Prosedur / Procedure / " + ZH_PROCEDURE,
        fontname="china-s",
        fontsize=17,
    )
    document.set_metadata(
        {
            "title": path.stem,
            "author": "Document Compliance System",
            "subject": "Synthetic Phase 8 fixture without company content",
        }
    )
    document.save(path, garbage=4, deflate=True)
    document.close()


def generate(output_root: Path) -> list[Path]:
    """Generate the fixture inventory required by the Phase 8 prompt."""
    directory = output_root / "compliance"
    directory.mkdir(parents=True, exist_ok=True)
    outputs = [
        directory / "complete-three-language.docx",
        directory / "missing-chinese.docx",
        directory / "wrong-language-order.docx",
        directory / "missing-purpose-section.docx",
        directory / "incomplete-table.docx",
        directory / "three-language-columns.xlsx",
        directory / "missing-language-cell.xlsx",
        directory / "languages-as-rows.xlsx",
        directory / "three-language-pdf.pdf",
        directory / "wrong-order-pdf.pdf",
        directory / "low-confidence-grouping.pdf",
    ]
    _complete_docx(outputs[0])
    _complete_docx(outputs[1], include_chinese=False)
    _complete_docx(outputs[2], wrong_order=True)
    _missing_purpose_docx(outputs[3])
    _incomplete_table_docx(outputs[4])
    _columns_xlsx(outputs[5])
    _columns_xlsx(outputs[6], missing_cell=True)
    _rows_xlsx(outputs[7])
    _pdf(outputs[8])
    _pdf(outputs[9], wrong_order=True)
    _pdf(outputs[10], low_confidence=True)
    return outputs


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate synthetic Phase 8 compliance fixtures.",
    )
    parser.add_argument(
        "--output-root",
        type=Path,
        default=Path(__file__).resolve().parents[2] / "sample-documents",
    )
    arguments = parser.parse_args()
    for output in generate(arguments.output_root.resolve()):
        print(output)


if __name__ == "__main__":
    main()
