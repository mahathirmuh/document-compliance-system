"""Generate synthetic SharePoint/synchronisation fixtures for Phase 10."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

OUTPUT_ROOT = (
    Path(__file__).resolve().parents[2] / "sample-documents" / "sharepoint"
)


def _write_pdf(path: Path, title: str, lines: list[str]) -> None:
    pdf = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    pdf.setTitle(title)
    pdf.setFont("Helvetica-Bold", 15)
    pdf.drawString(54, height - 64, title)
    pdf.setFont("Helvetica", 10)
    y = height - 96
    for line in lines:
        pdf.drawString(54, y, line)
        y -= 18
    pdf.drawString(54, 48, "Synthetic Phase 10 fixture - no company data")
    pdf.save()


def _write_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Inbound SharePoint Procedure", level=1)
    document.add_paragraph(
        "Kode Dokumen / Document Code / 文档代码: MOCK-SP-IN-001"
    )
    document.add_heading("Purpose / Tujuan / 目的", level=2)
    document.add_paragraph(
        "This generated file validates inbound metadata and content flow."
    )
    document.add_paragraph(
        "Dokumen sintetis ini menguji alur metadata dan konten masuk."
    )
    document.add_paragraph("此合成文件用于验证入站元数据与内容流程。")
    document.save(path)


def _write_workbook(path: Path, *, source: str, quantity: int) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "ControlledData"
    sheet.append(["DocumentCode", "Source", "Quantity", "Approved"])
    sheet.append(["MOCK-SP-XLSX-001", source, quantity, True])
    sheet.append(["MOCK-SP-XLSX-002", source, quantity + 1, False])
    workbook.save(path)


def main() -> None:
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    _write_pdf(
        OUTPUT_ROOT / "outbound-document.pdf",
        "Outbound SharePoint Document",
        [
            "Document Code: MOCK-SP-OUT-001",
            "Revision: 01",
            "Department: QA",
            "Status: Approved",
        ],
    )
    _write_docx(OUTPUT_ROOT / "inbound-document.docx")
    _write_workbook(
        OUTPUT_ROOT / "changed-local.xlsx",
        source="LOCAL",
        quantity=10,
    )
    _write_workbook(
        OUTPUT_ROOT / "changed-remote.xlsx",
        source="SHAREPOINT",
        quantity=12,
    )
    _write_pdf(
        OUTPUT_ROOT / "conflict-document.pdf",
        "Bidirectional Conflict Fixture",
        [
            "Document Code: MOCK-SP-CONFLICT-001",
            "Local eTag: local-etag-02",
            "Remote eTag: remote-etag-03",
            "Expected policy: MANUAL",
        ],
    )
    print(f"Generated 5 Phase 10 fixtures in {OUTPUT_ROOT}")


if __name__ == "__main__":
    main()
