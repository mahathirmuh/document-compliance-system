"""Generate synthetic Phase 7 OCR and language-detection fixtures."""

from __future__ import annotations

import argparse
from io import BytesIO
from pathlib import Path

import pymupdf
from docx import Document
from openpyxl import Workbook
from PIL import Image

PAGE_WIDTH = 595
PAGE_HEIGHT = 842


def _render_lines(
    lines: list[tuple[str, str]],
    *,
    dpi: int = 150,
) -> bytes:
    document = pymupdf.open()
    page = document.new_page(width=PAGE_WIDTH, height=PAGE_HEIGHT)
    y = 90
    for text, font_name in lines:
        page.insert_text(
            (55, y),
            text,
            fontname=font_name,
            fontsize=18,
            color=(0.04, 0.04, 0.04),
        )
        y += 42
    pixmap = page.get_pixmap(
        matrix=pymupdf.Matrix(dpi / 72, dpi / 72),
        colorspace=pymupdf.csRGB,
        alpha=False,
    )
    payload = pixmap.tobytes("png")
    document.close()
    return payload


def _rotate_image(payload: bytes, degrees: int) -> bytes:
    with Image.open(BytesIO(payload)) as source:
        rotated = source.rotate(
            degrees,
            expand=True,
            fillcolor="white",
        )
        output = BytesIO()
        rotated.save(output, format="PNG")
        return output.getvalue()


def _degrade_image(payload: bytes) -> bytes:
    with Image.open(BytesIO(payload)) as source:
        reduced = source.resize(
            (
                max(1, source.width // 5),
                max(1, source.height // 5),
            ),
        )
        output = BytesIO()
        reduced.save(output, format="JPEG", quality=28, optimize=True)
        return output.getvalue()


def _blank_image() -> bytes:
    output = BytesIO()
    Image.new("RGB", (420, 594), "white").save(
        output,
        format="JPEG",
        quality=80,
    )
    return output.getvalue()


def _insert_image_page(document: pymupdf.Document, payload: bytes) -> None:
    page = document.new_page(width=PAGE_WIDTH, height=PAGE_HEIGHT)
    page.insert_image(
        page.rect,
        stream=payload,
        keep_proportion=True,
    )


def _write_scanned_pdf(path: Path, pages: list[bytes]) -> None:
    document = pymupdf.open()
    for payload in pages:
        _insert_image_page(document, payload)
    document.set_metadata(
        {
            "title": f"Synthetic Phase 7 fixture: {path.stem}",
            "author": "Document Compliance System",
            "subject": "Generated test data without company content",
        }
    )
    document.save(path, garbage=4, deflate=True)
    document.close()


def _write_partial_scan(path: Path, scanned_page: bytes) -> None:
    document = pymupdf.open()
    native_page = document.new_page(width=PAGE_WIDTH, height=PAGE_HEIGHT)
    native_page.insert_textbox(
        pymupdf.Rect(55, 75, 540, 360),
        (
            "Native selectable page\n\n"
            "This page intentionally contains enough selectable text to be "
            "skipped by automatic OCR page selection. The following page is "
            "an image-only scan and must remain eligible for OCR."
        ),
        fontname="helv",
        fontsize=15,
    )
    _insert_image_page(document, scanned_page)
    document.set_metadata(
        {
            "title": "Synthetic Phase 7 partial-scan fixture",
            "author": "Document Compliance System",
            "subject": "Generated test data without company content",
        }
    )
    document.save(path, garbage=4, deflate=True)
    document.close()


def _write_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    document = Document()
    document.core_properties.author = "Document Compliance System"
    document.core_properties.title = title
    document.core_properties.subject = "Synthetic Phase 7 language fixture"
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)


def _write_xlsx(path: Path, rows: list[tuple[str, str]]) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Generated content"
    sheet.append(("Label", "Text"))
    for label, text in rows:
        sheet.append((label, text))
    sheet.freeze_panes = "A2"
    sheet.column_dimensions["A"].width = 22
    sheet.column_dimensions["B"].width = 100
    workbook.properties.creator = "Document Compliance System"
    workbook.properties.title = f"Synthetic Phase 7 fixture: {path.stem}"
    workbook.properties.subject = "Generated test data without company content"
    workbook.save(path)


def generate(output_root: Path) -> list[Path]:
    """Generate the exact fixture inventory required by Phase 7."""
    ocr_directory = output_root / "ocr"
    language_directory = output_root / "language"
    ocr_directory.mkdir(parents=True, exist_ok=True)
    language_directory.mkdir(parents=True, exist_ok=True)

    indonesian = _render_lines(
        [
            ("Dokumen ini dibuat khusus untuk pengujian OCR lokal.", "helv"),
            ("Kebijakan mutu harus ditinjau secara berkala.", "helv"),
            ("Setiap perubahan dicatat dan dapat ditelusuri.", "helv"),
            ("Tidak ada data perusahaan di dalam berkas ini.", "helv"),
        ]
    )
    english = _render_lines(
        [
            ("This document is generated for local OCR testing.", "helv"),
            ("Quality procedures are reviewed on a regular schedule.", "helv"),
            ("Every revision remains traceable in the document register.", "helv"),
            ("The fixture contains no internal company information.", "helv"),
        ]
    )
    chinese = _render_lines(
        [
            ("本文件仅用于本地文字识别测试。", "china-s"),
            ("质量管理程序应当定期审查。", "china-s"),
            ("每次修订都必须保留完整记录。", "china-s"),
            ("此测试文件不包含任何公司数据。", "china-s"),
        ]
    )
    mixed = _render_lines(
        [
            ("Dokumen pengujian OCR dalam tiga bahasa.", "helv"),
            ("This generated page validates mixed-language blocks.", "helv"),
            ("本页面用于验证多语言文字识别。", "china-s"),
            ("Semua konten bersifat sintetis dan aman.", "helv"),
        ]
    )

    outputs = [
        ocr_directory / "scanned-indonesian.pdf",
        ocr_directory / "scanned-english.pdf",
        ocr_directory / "scanned-chinese.pdf",
        ocr_directory / "scanned-mixed.pdf",
        ocr_directory / "rotated.pdf",
        ocr_directory / "low-resolution.pdf",
        ocr_directory / "partial-scan.pdf",
        ocr_directory / "blank-scan.pdf",
        language_directory / "indonesian.docx",
        language_directory / "english.docx",
        language_directory / "chinese.docx",
        language_directory / "three-language.docx",
        language_directory / "mixed-language.xlsx",
        language_directory / "short-text.xlsx",
    ]

    _write_scanned_pdf(outputs[0], [indonesian])
    _write_scanned_pdf(outputs[1], [english])
    _write_scanned_pdf(outputs[2], [chinese])
    _write_scanned_pdf(outputs[3], [mixed])
    _write_scanned_pdf(outputs[4], [_rotate_image(english, 90)])
    _write_scanned_pdf(outputs[5], [_degrade_image(indonesian)])
    _write_partial_scan(outputs[6], english)
    _write_scanned_pdf(outputs[7], [_blank_image()])

    _write_docx(
        outputs[8],
        "Dokumen Bahasa Indonesia",
        [
            "Dokumen sintetis ini digunakan untuk menguji deteksi bahasa Indonesia.",
            "Prosedur pengendalian dokumen memastikan setiap revisi dapat ditelusuri.",
        ],
    )
    _write_docx(
        outputs[9],
        "English Language Document",
        [
            "This synthetic document is used to test English language detection.",
            "Document control ensures that every approved revision remains traceable.",
        ],
    )
    _write_docx(
        outputs[10],
        "中文测试文件",
        [
            "这份合成文件用于测试中文语言检测。",
            "文件控制程序确保每个批准的修订版本都可以追溯。",
        ],
    )
    _write_docx(
        outputs[11],
        "Three-language Test Document",
        [
            "Dokumen sintetis ini memuat paragraf bahasa Indonesia untuk pengujian.",
            "This synthetic paragraph provides a clear English language sample.",
            "这个合成段落提供了清晰的中文语言样本。",
        ],
    )
    _write_xlsx(
        outputs[12],
        [
            ("Indonesian", "Prosedur ini menjelaskan pengendalian dokumen mutu."),
            ("English", "This procedure describes quality document control."),
            ("Chinese", "本程序说明质量文件的控制要求。"),
            ("Mixed", "Status disetujui, approved, 已批准。"),
        ],
    )
    _write_xlsx(
        outputs[13],
        [
            ("Code", "QMS-001"),
            ("Number", "2026"),
            ("Short Indonesian", "disetujui"),
            ("Short English", "approved"),
            ("Short Chinese", "批准"),
            ("URL", "https://example.invalid/test"),
        ],
    )
    return outputs


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate synthetic Phase 7 sample documents.",
    )
    parser.add_argument(
        "--output-root",
        type=Path,
        default=Path(__file__).resolve().parents[2] / "sample-documents",
    )
    arguments = parser.parse_args()
    outputs = generate(arguments.output_root.resolve())
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
