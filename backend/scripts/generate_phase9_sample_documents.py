"""Generate deterministic, synthetic Phase 9 quality-intelligence fixtures."""

from __future__ import annotations

import argparse
import re
from datetime import UTC, datetime
from pathlib import Path
from typing import Final
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.document import Document as DocumentType
from openpyxl import Workbook
from openpyxl.styles import Font
from openpyxl.worksheet.worksheet import Worksheet

FIXED_DOCUMENT_TIME: Final = datetime(2026, 1, 1, 0, 0, 0, tzinfo=UTC)
FIXED_ZIP_TIME: Final = (2026, 1, 1, 0, 0, 0)
FIXTURE_AUTHOR: Final = "Document Compliance System"
FIXTURE_SUBJECT: Final = "Synthetic Phase 9 fixture without company content"
CORE_MODIFIED_PATTERN: Final = re.compile(
    rb"(<dcterms:modified\b[^>]*>).*?(</dcterms:modified>)"
)
FIXED_CORE_MODIFIED: Final = rb"\g<1>2026-01-01T00:00:00Z\g<2>"

EXPECTED_RELATIVE_PATHS: Final = (
    Path("similarity/equivalent-three-language.docx"),
    Path("similarity/unrelated-chinese-translation.docx"),
    Path("similarity/number-mismatch.docx"),
    Path("similarity/date-mismatch.docx"),
    Path("similarity/measurement-mismatch.docx"),
    Path("similarity/reference-mismatch.docx"),
    Path("similarity/negation-mismatch.docx"),
    Path("glossary/preferred-terms.docx"),
    Path("glossary/forbidden-terms.docx"),
    Path("glossary/missing-translation.docx"),
    Path("glossary/inconsistent-terms.docx"),
    Path("glossary/chinese-terms.xlsx"),
    Path("revisions/revision-001.docx"),
    Path("revisions/revision-002.docx"),
    Path("revisions/revision-001.xlsx"),
    Path("revisions/revision-002.xlsx"),
)

TrilingualRow = tuple[str, str, str, str]


def _normalize_ooxml_archive(path: Path) -> None:
    """Normalize ZIP metadata so repeated generations are byte-identical."""
    temporary_path = path.with_suffix(f"{path.suffix}.normalized")
    with ZipFile(path, "r") as source:
        entries = [
            (
                entry.filename,
                (
                    CORE_MODIFIED_PATTERN.sub(
                        FIXED_CORE_MODIFIED,
                        source.read(entry.filename),
                    )
                    if entry.filename == "docProps/core.xml"
                    else source.read(entry.filename)
                ),
            )
            for entry in sorted(source.infolist(), key=lambda item: item.filename)
        ]

    with ZipFile(
        temporary_path,
        "w",
        compression=ZIP_DEFLATED,
        compresslevel=9,
    ) as target:
        for filename, payload in entries:
            entry = ZipInfo(filename=filename, date_time=FIXED_ZIP_TIME)
            entry.compress_type = ZIP_DEFLATED
            entry.create_system = 0
            entry.external_attr = 0
            target.writestr(entry, payload, compress_type=ZIP_DEFLATED, compresslevel=9)

    temporary_path.replace(path)


def _set_document_metadata(document: DocumentType, title: str) -> None:
    properties = document.core_properties
    properties.author = FIXTURE_AUTHOR
    properties.last_modified_by = FIXTURE_AUTHOR
    properties.title = title
    properties.subject = FIXTURE_SUBJECT
    properties.comments = "Generated deterministically for local automated testing."
    properties.created = FIXED_DOCUMENT_TIME
    properties.modified = FIXED_DOCUMENT_TIME
    properties.revision = 9


def _add_trilingual_table(
    document: DocumentType,
    rows: tuple[TrilingualRow, ...],
) -> None:
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ("Item", "Bahasa Indonesia", "English", "\u4e2d\u6587")
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        cell.text = value
    for item, indonesian, english, chinese in rows:
        cells = table.add_row().cells
        for cell, value in zip(
            cells,
            (item, indonesian, english, chinese),
            strict=True,
        ):
            cell.text = value


def _save_docx(
    path: Path,
    *,
    title: str,
    rows: tuple[TrilingualRow, ...],
    note: str,
) -> None:
    document = Document()
    _set_document_metadata(document, title)
    document.add_heading(title, level=0)
    document.add_paragraph(note)
    document.add_heading("Prosedur / Procedure / \u7a0b\u5e8f", level=1)
    _add_trilingual_table(document, rows)
    document.save(path)
    _normalize_ooxml_archive(path)


def _set_workbook_metadata(workbook: Workbook, title: str) -> None:
    workbook.properties.creator = FIXTURE_AUTHOR
    workbook.properties.lastModifiedBy = FIXTURE_AUTHOR
    workbook.properties.title = title
    workbook.properties.subject = FIXTURE_SUBJECT
    workbook.properties.description = (
        "Generated deterministically for local automated testing."
    )
    workbook.properties.created = FIXED_DOCUMENT_TIME.replace(tzinfo=None)
    workbook.properties.modified = FIXED_DOCUMENT_TIME.replace(tzinfo=None)


def _style_sheet(sheet: Worksheet) -> None:
    for cell in sheet[1]:
        cell.font = Font(bold=True)
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    for column in sheet.columns:
        letter = column[0].column_letter
        longest = max(len(str(cell.value or "")) for cell in column)
        sheet.column_dimensions[letter].width = min(max(longest + 2, 12), 48)


def _save_workbook(workbook: Workbook, path: Path) -> None:
    workbook.calculation.fullCalcOnLoad = False
    workbook.calculation.forceFullCalc = False
    workbook.calculation.calcMode = "manual"
    workbook.save(path)
    _normalize_ooxml_archive(path)


def _generate_similarity_fixtures(directory: Path) -> list[Path]:
    directory.mkdir(parents=True, exist_ok=True)
    fixtures: tuple[tuple[str, str, tuple[TrilingualRow, ...]], ...] = (
        (
            "equivalent-three-language.docx",
            "Equivalent three-language synthetic procedure",
            (
                (
                    "SIM-001",
                    "Petugas memeriksa catatan uji sebelum dokumen diterbitkan.",
                    "The officer checks the test record before the document is issued.",
                    (
                        "\u4eba\u5458\u5728\u6587\u4ef6\u53d1\u5e03\u524d"
                        "\u68c0\u67e5\u6d4b\u8bd5\u8bb0\u5f55\u3002"
                    ),
                ),
                (
                    "SIM-002",
                    "Pemilik dokumen menyimpan bukti tinjauan pada arsip lokal.",
                    "The document owner retains review evidence in the local archive.",
                    (
                        "\u6587\u4ef6\u8d1f\u8d23\u4eba\u5c06\u5ba1\u6838"
                        "\u8bc1\u636e\u4fdd\u7559\u5728\u672c\u5730\u6863\u6848\u4e2d\u3002"
                    ),
                ),
            ),
        ),
        (
            "unrelated-chinese-translation.docx",
            "Unrelated Chinese translation signal",
            (
                (
                    "SIM-003",
                    "Petugas mengunci lemari arsip setelah pemeriksaan selesai.",
                    "The officer locks the archive cabinet after the inspection is complete.",
                    (
                        "\u4eca\u5929\u5929\u6c14\u6674\u6717\uff0c"
                        "\u82b1\u56ed\u91cc\u6709\u5f88\u591a\u9e1f\u3002"
                    ),
                ),
            ),
        ),
        (
            "number-mismatch.docx",
            "Numeric consistency mismatch",
            (
                (
                    "SIM-004",
                    "Tim meninjau 24 catatan uji setiap bulan.",
                    "The team reviews 25 test records every month.",
                    (
                        "\u56e2\u961f\u6bcf\u6708\u5ba1\u67e5"
                        "24\u4efd\u6d4b\u8bd5\u8bb0\u5f55\u3002"
                    ),
                ),
            ),
        ),
        (
            "date-mismatch.docx",
            "Date consistency mismatch",
            (
                (
                    "SIM-005",
                    "Tinjauan berikutnya dijadwalkan pada 15 Agustus 2026.",
                    "The next review is scheduled for 16 August 2026.",
                    (
                        "\u4e0b\u6b21\u5ba1\u67e5\u5b9a\u4e8e"
                        "2026\u5e748\u670815\u65e5\u3002"
                    ),
                ),
            ),
        ),
        (
            "measurement-mismatch.docx",
            "Measurement consistency mismatch",
            (
                (
                    "SIM-006",
                    "Sampel uji harus memiliki massa 5 kg.",
                    "The test sample must have a mass of 7 kg.",
                    (
                        "\u6d4b\u8bd5\u6837\u54c1\u7684\u8d28\u91cf"
                        "\u5fc5\u987b\u4e3a5\u5343\u514b\u3002"
                    ),
                ),
            ),
        ),
        (
            "reference-mismatch.docx",
            "Reference consistency mismatch",
            (
                (
                    "SIM-007",
                    "Gunakan instruksi kerja WI-TEST-009 untuk pemeriksaan.",
                    "Use work instruction WI-TEST-010 for the inspection.",
                    (
                        "\u68c0\u67e5\u65f6\u4f7f\u7528\u4f5c\u4e1a"
                        "\u6307\u5bfc\u4e66WI-TEST-009\u3002"
                    ),
                ),
            ),
        ),
        (
            "negation-mismatch.docx",
            "Negation consistency review signal",
            (
                (
                    "SIM-008",
                    "Operator tidak boleh menghapus catatan uji yang telah disetujui.",
                    "The operator may delete an approved test record.",
                    (
                        "\u64cd\u4f5c\u5458\u4e0d\u5f97\u5220\u9664"
                        "\u5df2\u6279\u51c6\u7684\u6d4b\u8bd5\u8bb0\u5f55\u3002"
                    ),
                ),
            ),
        ),
    )
    outputs: list[Path] = []
    for filename, title, rows in fixtures:
        path = directory / filename
        _save_docx(
            path,
            title=title,
            rows=rows,
            note=(
                "Synthetic content for local translation-similarity and "
                "consistency testing."
            ),
        )
        outputs.append(path)
    return outputs


def _generate_glossary_workbook(path: Path) -> None:
    workbook = Workbook()
    _set_workbook_metadata(workbook, "Synthetic Chinese glossary import")

    terms = workbook.active
    terms.title = "Terms"
    terms.append(
        (
            "profile_code",
            "term_code",
            "concept_name",
            "description",
            "term_type",
            "severity",
            "is_case_sensitive",
            "match_whole_word",
            "allow_inflection",
            "is_regex",
            "is_active",
            "notes",
        )
    )
    terms.append(
        (
            "PHASE9_SAMPLE",
            "CONTROLLED_DOCUMENT",
            "Controlled document",
            "A synthetic document whose revisions are controlled.",
            "PREFERRED",
            "MAJOR",
            False,
            True,
            False,
            False,
            True,
            "Synthetic sample only",
        )
    )

    translations = workbook.create_sheet("Translations")
    translations.append(
        (
            "term_code",
            "language_code",
            "term_text",
            "is_preferred",
            "is_forbidden",
            "is_required",
            "priority",
            "usage_note",
            "example_text",
            "is_active",
        )
    )
    translations.append(
        (
            "CONTROLLED_DOCUMENT",
            "id",
            "dokumen terkendali",
            True,
            False,
            True,
            1,
            "Gunakan istilah pilihan.",
            "Simpan dokumen terkendali.",
            True,
        )
    )
    translations.append(
        (
            "CONTROLLED_DOCUMENT",
            "en",
            "controlled document",
            True,
            False,
            True,
            1,
            "Use the preferred term.",
            "Retain the controlled document.",
            True,
        )
    )
    translations.append(
        (
            "CONTROLLED_DOCUMENT",
            "zh",
            "\u53d7\u63a7\u6587\u4ef6",
            True,
            False,
            True,
            1,
            "\u4f7f\u7528\u9996\u9009\u672f\u8bed\u3002",
            "\u4fdd\u7559\u53d7\u63a7\u6587\u4ef6\u3002",
            True,
        )
    )

    variants = workbook.create_sheet("Variants")
    variants.append(
        (
            "term_code",
            "language_code",
            "preferred_term_text",
            "variant_text",
            "variant_type",
            "is_allowed",
            "is_active",
        )
    )
    variants.append(
        (
            "CONTROLLED_DOCUMENT",
            "zh",
            "\u53d7\u63a7\u6587\u4ef6",
            "\u7ba1\u5236\u6587\u6863",
            "LEGACY",
            False,
            True,
        )
    )

    for sheet in workbook.worksheets:
        _style_sheet(sheet)
    _save_workbook(workbook, path)


def _generate_glossary_fixtures(directory: Path) -> list[Path]:
    directory.mkdir(parents=True, exist_ok=True)
    fixtures: tuple[tuple[str, str, tuple[TrilingualRow, ...]], ...] = (
        (
            "preferred-terms.docx",
            "Preferred glossary terms",
            (
                (
                    "GLO-001",
                    "Simpan dokumen terkendali di dalam arsip lokal.",
                    "Retain the controlled document in the local archive.",
                    (
                        "\u5c06\u53d7\u63a7\u6587\u4ef6"
                        "\u4fdd\u7559\u5728\u672c\u5730\u6863\u6848\u4e2d\u3002"
                    ),
                ),
            ),
        ),
        (
            "forbidden-terms.docx",
            "Forbidden glossary terms",
            (
                (
                    "GLO-002",
                    "Simpan dokumen master lama untuk pemeriksaan.",
                    "Retain the legacy master copy for inspection.",
                    (
                        "\u4fdd\u7559\u65e7\u7684\u4e3b\u6587\u6863"
                        "\u4ee5\u4f9b\u68c0\u67e5\u3002"
                    ),
                ),
            ),
        ),
        (
            "missing-translation.docx",
            "Missing glossary translation",
            (
                (
                    "GLO-003",
                    "Pemilik menyimpan rekaman uji setelah peninjauan.",
                    "The owner retains the test record after review.",
                    (
                        "\u8d1f\u8d23\u4eba\u5728\u5ba1\u6838"
                        "\u540e\u4fdd\u7559\u8d44\u6599\u3002"
                    ),
                ),
            ),
        ),
        (
            "inconsistent-terms.docx",
            "Inconsistent glossary terms",
            (
                (
                    "GLO-004-A",
                    "Simpan dokumen terkendali setelah persetujuan.",
                    "Retain the controlled document after approval.",
                    "\u6279\u51c6\u540e\u4fdd\u7559\u53d7\u63a7\u6587\u4ef6\u3002",
                ),
                (
                    "GLO-004-B",
                    "Tinjau dokumen terkelola setiap tahun.",
                    "Review the managed document every year.",
                    "\u6bcf\u5e74\u5ba1\u67e5\u7ba1\u5236\u6587\u6863\u3002",
                ),
            ),
        ),
    )
    outputs: list[Path] = []
    for filename, title, rows in fixtures:
        path = directory / filename
        _save_docx(
            path,
            title=title,
            rows=rows,
            note="Synthetic content for local glossary-validation testing.",
        )
        outputs.append(path)

    workbook_path = directory / "chinese-terms.xlsx"
    _generate_glossary_workbook(workbook_path)
    outputs.append(workbook_path)
    return outputs


def _revision_docx_rows(revision: int) -> tuple[TrilingualRow, ...]:
    if revision == 1:
        return (
            (
                "REV-ITEM-001",
                "Satu peninjau memeriksa catatan uji sebelum penerbitan.",
                "One reviewer checks the test record before issuance.",
                (
                    "\u4e00\u540d\u5ba1\u6838\u5458\u5728\u53d1\u5e03"
                    "\u524d\u68c0\u67e5\u6d4b\u8bd5\u8bb0\u5f55\u3002"
                ),
            ),
            (
                "REV-ITEM-002",
                "Bukti tinjauan disimpan selama 12 bulan.",
                "Review evidence is retained for 12 months.",
                ("\u5ba1\u6838\u8bc1\u636e\u4fdd\u755912\u4e2a\u6708\u3002"),
            ),
            (
                "REV-ITEM-003",
                "Salinan usang dipindahkan ke arsip lokal.",
                "Obsolete copies are moved to the local archive.",
                ("\u4f5c\u5e9f\u526f\u672c\u79fb\u81f3\u672c\u5730\u6863\u6848\u3002"),
            ),
        )
    return (
        (
            "REV-ITEM-001",
            "Dua peninjau memeriksa catatan uji sebelum penerbitan.",
            "Two reviewers check the test record before issuance.",
            (
                "\u4e24\u540d\u5ba1\u6838\u5458\u5728\u53d1\u5e03"
                "\u524d\u68c0\u67e5\u6d4b\u8bd5\u8bb0\u5f55\u3002"
            ),
        ),
        (
            "REV-ITEM-003",
            "Salinan usang dipindahkan ke arsip lokal.",
            "Obsolete copies are moved to the local archive.",
            ("\u4f5c\u5e9f\u526f\u672c\u79fb\u81f3\u672c\u5730\u6863\u6848\u3002"),
        ),
        (
            "REV-ITEM-004",
            "Pemilik mencatat alasan perubahan pada log revisi.",
            "The owner records the change reason in the revision log.",
            (
                "\u8d1f\u8d23\u4eba\u5728\u4fee\u8ba2"
                "\u65e5\u5fd7\u4e2d\u8bb0\u5f55\u53d8\u66f4\u539f\u56e0\u3002"
            ),
        ),
    )


def _generate_revision_workbook(path: Path, *, revision: int) -> None:
    workbook = Workbook()
    _set_workbook_metadata(workbook, f"Synthetic revision {revision:03d}")
    sheet = workbook.active
    sheet.title = "Revision Items"
    sheet.append(
        (
            "Item",
            "Bahasa Indonesia",
            "English",
            "\u4e2d\u6587",
            "Owner",
            "Retention Months",
        )
    )
    if revision == 1:
        rows = (
            (
                "REV-XLSX-001",
                "Periksa catatan uji.",
                "Check the test record.",
                "\u68c0\u67e5\u6d4b\u8bd5\u8bb0\u5f55\u3002",
                "Reviewer",
                12,
            ),
            (
                "REV-XLSX-002",
                "Simpan bukti tinjauan.",
                "Retain review evidence.",
                "\u4fdd\u7559\u5ba1\u6838\u8bc1\u636e\u3002",
                "Owner",
                12,
            ),
            (
                "REV-XLSX-003",
                "Arsipkan salinan usang.",
                "Archive obsolete copies.",
                "\u5f52\u6863\u4f5c\u5e9f\u526f\u672c\u3002",
                "Controller",
                24,
            ),
        )
    else:
        rows = (
            (
                "REV-XLSX-001",
                "Dua peninjau memeriksa catatan uji.",
                "Two reviewers check the test record.",
                (
                    "\u4e24\u540d\u5ba1\u6838\u5458"
                    "\u68c0\u67e5\u6d4b\u8bd5\u8bb0\u5f55\u3002"
                ),
                "Senior Reviewer",
                18,
            ),
            (
                "REV-XLSX-003",
                "Arsipkan salinan usang.",
                "Archive obsolete copies.",
                "\u5f52\u6863\u4f5c\u5e9f\u526f\u672c\u3002",
                "Controller",
                24,
            ),
            (
                "REV-XLSX-004",
                "Catat alasan perubahan.",
                "Record the change reason.",
                "\u8bb0\u5f55\u53d8\u66f4\u539f\u56e0\u3002",
                "Owner",
                36,
            ),
        )
    for row in rows:
        sheet.append(row)
    _style_sheet(sheet)
    _save_workbook(workbook, path)


def _generate_revision_fixtures(directory: Path) -> list[Path]:
    directory.mkdir(parents=True, exist_ok=True)
    outputs: list[Path] = []
    for revision in (1, 2):
        docx_path = directory / f"revision-{revision:03d}.docx"
        _save_docx(
            docx_path,
            title=f"Synthetic procedure revision {revision:03d}",
            rows=_revision_docx_rows(revision),
            note=(
                "Synthetic revision pair with added, removed, modified, "
                "and unchanged items."
            ),
        )
        outputs.append(docx_path)
    for revision in (1, 2):
        xlsx_path = directory / f"revision-{revision:03d}.xlsx"
        _generate_revision_workbook(xlsx_path, revision=revision)
        outputs.append(xlsx_path)
    return outputs


def generate(output_root: Path) -> list[Path]:
    """Generate the exact fixture inventory required by the Phase 9 prompt."""
    root = output_root.resolve()
    outputs = [
        *_generate_similarity_fixtures(root / "similarity"),
        *_generate_glossary_fixtures(root / "glossary"),
        *_generate_revision_fixtures(root / "revisions"),
    ]
    actual_paths = tuple(output.relative_to(root) for output in outputs)
    if actual_paths != EXPECTED_RELATIVE_PATHS:
        raise RuntimeError("Generated Phase 9 fixture inventory is not canonical.")
    return outputs


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate deterministic synthetic Phase 9 fixtures.",
    )
    parser.add_argument(
        "--output-root",
        type=Path,
        default=Path(__file__).resolve().parents[2] / "sample-documents",
    )
    arguments = parser.parse_args()
    for output in generate(arguments.output_root):
        print(output)


if __name__ == "__main__":
    main()
